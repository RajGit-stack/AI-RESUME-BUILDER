import tempfile
import os
import markdown
import html2text
from typing import Optional, Dict
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_CENTER
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.services.template_html_generator import generate_template_html

# Try to import PDF generation libraries (optional)
try:
    from xhtml2pdf import pisa
    HAS_XHTML2PDF = True
except ImportError:
    HAS_XHTML2PDF = False

try:
    from weasyprint import HTML
    HAS_WEASYPRINT = True
except (ImportError, OSError):
    HAS_WEASYPRINT = False

def generate_pdf(content: str, title: str, template: Optional[Dict] = None, personal_info: Optional[Dict] = None, customization: Optional[Dict] = None) -> str:
    """Generate PDF from markdown content with template styling"""
    # Create temporary file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
    temp_file.close()
    
    # Ensure template is a dict with 'id' if it's not None
    if template is None:
        template = {'id': 'minimalist-clean', 'name': 'Minimalist Clean'}
    elif isinstance(template, dict) and 'id' not in template:
        template['id'] = template.get('id', 'minimalist-clean')
    
    print(f"[PDF DEBUG] Template: {template}, Type: {type(template)}")
    
    # Generate HTML from template with one-page constraint
    html_content = generate_template_html(content, template, personal_info, customization, one_page=True)
    
    # Try different PDF generation methods
    pdf_generated = False
    
    # Method 1: Try WeasyPrint (best quality, requires system libs on Windows)
    if HAS_WEASYPRINT:
        try:
            HTML(string=html_content).write_pdf(temp_file.name)
            pdf_generated = True
        except Exception as e:
            print(f"WeasyPrint failed: {e}. Trying alternative...")
    
    # Method 2: Try xhtml2pdf (pure Python, works on Windows)
    if not pdf_generated and HAS_XHTML2PDF:
        try:
            from xhtml2pdf.config import httpConfig
            httpConfig.save_keys = False
            
            # Debug: Save HTML to file for inspection
            import os
            debug_html_file = temp_file.name.replace('.pdf', '_debug.html')
            with open(debug_html_file, 'w', encoding='utf-8') as f:
                f.write(html_content)
            print(f"Debug: HTML saved to {debug_html_file}")
            
            with open(temp_file.name, 'wb') as pdf_file:
                pisa_status = pisa.CreatePDF(
                    html_content,
                    dest=pdf_file,
                    encoding='utf-8',
                    link_callback=None
                )
            if not pisa_status.err:
                pdf_generated = True
                print("PDF generated successfully with xhtml2pdf")
            else:
                print(f"xhtml2pdf errors: {pisa_status.err}")
                # Try with alternative settings
                try:
                    with open(temp_file.name, 'wb') as pdf_file:
                        from xhtml2pdf import pisa
                        from io import BytesIO
                        result = BytesIO()
                        pisa_status = pisa.CreatePDF(
                            BytesIO(html_content.encode('utf-8')),
                            dest=result
                        )
                        if not pisa_status.err:
                            pdf_file.write(result.getvalue())
                            pdf_generated = True
                            print("PDF generated with alternative method")
                except Exception as e2:
                    print(f"Alternative method failed: {e2}")
        except Exception as e:
            print(f"xhtml2pdf exception: {e}")
            import traceback
            traceback.print_exc()
    
    # Method 3: Fallback to ReportLab with basic styling
    if not pdf_generated:
        print("Falling back to ReportLab plain PDF...")
        doc = SimpleDocTemplate(temp_file.name, pagesize=letter)
        story = []
        
        # Define styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor='#1e293b',
            spaceAfter=12,
            alignment=TA_CENTER
        )
        
        h = html2text.HTML2Text()
        h.ignore_links = True
        plain_text = h.handle(content)
        
        # Extract name for title if available
        from app.services.template_html_generator import extract_contact_info
        contact_info = extract_contact_info(content, personal_info)
        name = contact_info.get('name', title)
        
        # Add title
        p = Paragraph(name, title_style)
        story.append(p)
        story.append(Spacer(1, 8))
        
        # Add contact info
        contact_text = ''
        if contact_info.get('email'):
            contact_text += contact_info['email']
        if contact_info.get('phone'):
            if contact_text:
                contact_text += ' | '
            contact_text += contact_info['phone']
        if contact_info.get('location'):
            if contact_text:
                contact_text += ' | '
            contact_text += contact_info['location']
        
        if contact_text:
            contact_para = Paragraph(contact_text, styles['Normal'])
            contact_para.alignment = TA_CENTER
            story.append(contact_para)
            story.append(Spacer(1, 12))
        
        # Process content
        for line in plain_text.split('\n'):
            line = line.strip()
            if not line:
                story.append(Spacer(1, 6))
                continue
            
            # Skip contact info already added
            if '**Email:**' in line or '**Phone:**' in line or '**Location:**' in line:
                continue
            
            # Check for headers
            if line.startswith('# '):
                p = Paragraph(line[2:], styles['Heading1'])
                story.append(p)
                story.append(Spacer(1, 12))
            elif line.startswith('## '):
                p = Paragraph(line[3:], styles['Heading2'])
                story.append(p)
                story.append(Spacer(1, 8))
            elif line.startswith('### '):
                p = Paragraph(line[4:], styles['Heading3'])
                story.append(p)
                story.append(Spacer(1, 6))
            else:
                # Regular paragraph
                p = Paragraph(line, styles['Normal'])
                story.append(p)
                story.append(Spacer(1, 6))
        
        # Build PDF
        doc.build(story)
    
    return temp_file.name

def generate_docx(content: str, title: str, template: Optional[Dict] = None, personal_info: Optional[Dict] = None, customization: Optional[Dict] = None) -> str:
    """Generate DOCX from markdown content with template styling"""
    # Create temporary file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    temp_file.close()
    
    # Create Document
    doc = Document()
    
    # Get template styling info
    template_id = 'minimalist-clean'
    if template and isinstance(template, dict):
        template_id = template.get('id', 'minimalist-clean')
    
    # Set default font based on template
    style = doc.styles['Normal']
    font = style.font
    if template_id == 'professional-classic':
        font.name = 'Calibri'
    else:
        font.name = 'Arial'
    font.size = Pt(11)
    
    # Extract contact info
    from app.services.template_html_generator import extract_contact_info
    contact_info = extract_contact_info(content, personal_info)
    
    # Add header with contact info
    if template_id == 'professional-classic':
        # Professional Classic Header
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(contact_info.get('name', title))
        name_run.font.size = Pt(28)
        name_run.font.bold = True
        name_run.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)
        
        contact_para = doc.add_paragraph()
        if contact_info.get('email'):
            contact_para.add_run(contact_info['email'])
        if contact_info.get('phone'):
            if contact_info.get('email'):
                contact_para.add_run(' | ')
            contact_para.add_run(contact_info['phone'])
        if contact_info.get('location'):
            contact_para.add_run(' | ')
            contact_para.add_run(contact_info['location'])
        contact_para.runs[0].font.size = Pt(10)
        
        doc.add_paragraph()  # Spacing
    else:
        # Minimalist Clean Header
        name_para = doc.add_paragraph(alignment=WD_ALIGN_PARAGRAPH.CENTER)
        name_run = name_para.add_run(contact_info.get('name', title))
        name_run.font.size = Pt(24)
        name_run.font.bold = True
        
        contact_para = doc.add_paragraph(alignment=WD_ALIGN_PARAGRAPH.CENTER)
        if contact_info.get('email'):
            contact_para.add_run(contact_info['email'])
        if contact_info.get('phone'):
            if contact_info.get('email'):
                contact_para.add_run(' | ')
            contact_para.add_run(contact_info['phone'])
        if contact_info.get('location'):
            contact_para.add_run(' | ')
            contact_para.add_run(contact_info['location'])
        contact_para.runs[0].font.size = Pt(9.5)
        
        doc.add_paragraph()  # Spacing
    
    # Convert markdown to HTML
    html_content = markdown.markdown(content)
    h = html2text.HTML2Text()
    h.ignore_links = True
    plain_text = h.handle(html_content)
    
    # Process content
    for line in plain_text.split('\n'):
        line = line.strip()
        if not line:
            continue
        
        # Skip contact info already added
        if '**Email:**' in line or '**Phone:**' in line or '**Location:**' in line:
            continue
        
        # Check for headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
            if template_id == 'professional-classic':
                p.runs[0].font.color.rgb = RGBColor(0x25, 0x63, 0xeb)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
            if template_id == 'professional-classic':
                p.runs[0].font.color.rgb = RGBColor(0x25, 0x63, 0xeb)
                p.runs[0].font.size = Pt(14)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
        else:
            # Regular paragraph
            if line.startswith('- '):
                # Bullet point
                p = doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith('* '):
                p = doc.add_paragraph(line[2:], style='List Bullet')
            else:
                p = doc.add_paragraph(line)
    
    # Save document
    doc.save(temp_file.name)
    
    return temp_file.name

