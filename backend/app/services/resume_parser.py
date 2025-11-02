"""
Service for parsing uploaded resume files (PDF and DOCX)
Converts them to markdown format for editing
"""
import io
from typing import Optional
from docx import Document
import PyPDF2
import pdfplumber


def parse_docx(file_content: bytes) -> str:
    """
    Parse DOCX file and convert to markdown
    """
    try:
        doc = Document(io.BytesIO(file_content))
        markdown_content = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
                
            # Check for heading styles
            if para.style.name.startswith('Heading'):
                level = para.style.name.replace('Heading', '').strip()
                if level == '1':
                    markdown_content.append(f"# {text}\n")
                elif level == '2':
                    markdown_content.append(f"## {text}\n")
                elif level == '3':
                    markdown_content.append(f"### {text}\n")
                else:
                    markdown_content.append(f"#### {text}\n")
            else:
                # Regular paragraph
                markdown_content.append(f"{text}\n\n")
        
        # Parse tables
        for table in doc.tables:
            markdown_content.append("\n")  # Add spacing before table
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                markdown_content.append(f"| {row_text} |\n")
            markdown_content.append("\n")
        
        return "".join(markdown_content).strip()
    
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX file: {str(e)}")


def parse_pdf(file_content: bytes) -> str:
    """
    Parse PDF file and convert to markdown
    Tries pdfplumber first (better), falls back to PyPDF2
    """
    try:
        # Try pdfplumber first (better for formatting)
        return parse_pdf_pdfplumber(file_content)
    except Exception:
        # Fallback to PyPDF2
        try:
            return parse_pdf_pypdf2(file_content)
        except Exception as e:
            raise ValueError(f"Failed to parse PDF file: {str(e)}")


def parse_pdf_pdfplumber(file_content: bytes) -> str:
    """Parse PDF using pdfplumber (better formatting)"""
    markdown_content = []
    
    with pdfplumber.open(io.BytesIO(file_content)) as pdf:
        for page_num, page in enumerate(pdf.pages):
            text = page.extract_text()
            if text:
                # Clean up the text
                lines = text.split('\n')
                cleaned_lines = []
                
                for line in lines:
                    line = line.strip()
                    if not line:
                        continue
                    
                    # Detect potential headers (all caps, short lines, bold-like formatting)
                    if len(line) < 50 and line.isupper():
                        cleaned_lines.append(f"## {line.title()}\n")
                    elif len(line) < 80 and not any(c.islower() for c in line):
                        cleaned_lines.append(f"### {line}\n")
                    else:
                        cleaned_lines.append(f"{line}\n")
                
                markdown_content.extend(cleaned_lines)
                markdown_content.append("\n")
    
    result = "\n".join(markdown_content).strip()
    return result if result else ""


def parse_pdf_pypdf2(file_content: bytes) -> str:
    """Parse PDF using PyPDF2 (fallback)"""
    markdown_content = []
    pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
    
    for page_num, page in enumerate(pdf_reader.pages):
        text = page.extract_text()
        if text:
            # Clean and format the text
            lines = text.split('\n')
            for line in lines:
                line = line.strip()
                if line:
                    markdown_content.append(f"{line}\n")
            markdown_content.append("\n")
    
    result = "\n".join(markdown_content).strip()
    return result if result else ""


def parse_resume_file(file_content: bytes, file_extension: str) -> str:
    """
    Main parsing function - routes to appropriate parser based on file type
    """
    file_extension = file_extension.lower().strip('.')
    
    if file_extension == 'docx':
        return parse_docx(file_content)
    elif file_extension == 'pdf':
        return parse_pdf(file_content)
    elif file_extension == 'doc':
        # Old DOC format - try to handle if possible
        raise ValueError("DOC format is not supported. Please convert to DOCX or PDF first.")
    else:
        raise ValueError(f"Unsupported file format: {file_extension}. Supported formats: PDF, DOCX")


